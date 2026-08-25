from __future__ import annotations

import random
from base64 import b64decode
from io import BytesIO
from itertools import product
from zipfile import ZIP_DEFLATED, BadZipFile, ZipFile

import pytest

from lumberjack.models import DocTree, DocumentBlock, SectionNode
from lumberjack.parser.code import NotebookParser, SourceCodeParser, SQLParser
from lumberjack.parser.docx import DocxParser
from lumberjack.parser.html import HTMLParser
from lumberjack.parser.markdown import MarkdownParser
from lumberjack.parser.records import (
    DelimitedTextParser,
    JSONLinesParser,
    JSONParser,
    TOMLParser,
    XMLParser,
    YAMLParser,
)
from lumberjack.parser.sqlite import SQLiteParser
from lumberjack.parser.xlsx import XlsxParser

MARKDOWN_BLOCK_CASES = {
    "paragraph": "{inline}\n",
    "heading": "# Heading\n\n{inline}\n",
    "heading-gap": "# Parent\n\n#### Deep\n\n{inline}\n",
    "blockquote": "> {inline}\n",
    "nested-blockquote": "> outer\n>\n> > {inline}\n",
    "bullet-list": "- {inline}\n- second item\n",
    "ordered-list": "7. {inline}\n8. second item\n",
    "nested-list": "- parent\n  1. {inline}\n  2. second item\n",
    "loose-list": "- {inline}\n\n- second item\n",
    "fence": "```text\n{inline}\n```\n",
    "tilde-fence": "~~~~ info with spaces\n{inline}\n~~~~\n",
    "indented-code": "    {inline}\n",
    "table": "| key | value |\n| --- | --- |\n| marker | {inline} |\n",
    "html": "<div>\n{inline}\n</div>\n",
    "math": "$$\n{inline}\n$$\n",
    "bracket-math": "\\[\n{inline}\n\\]\n",
    "setext-disabled": "Setext-like title\n=====\n\n{inline}\n",
    "reference": '{inline} [ref]\n\n[ref]: https://example.com "title"\n',
    "front-matter": "---\ntitle: Robustness\ntags: [one, two]\n---\n\n{inline}\n",
    "mixed-newlines": "# Heading\r\n\r\n{inline}\r\n",
}

MARKDOWN_INLINE_CASES = {
    "plain": "LJ_SENTINEL plain text",
    "unicode": "LJ_SENTINEL 中文 Ελληνικά العربية 👩🏽‍💻",
    "emphasis": "LJ_SENTINEL *em* **strong** ***both***",
    "underscores": "LJ_SENTINEL _em_ __strong__ snake_case",
    "code": "LJ_SENTINEL ``code with ` tick``",
    "link": 'LJ_SENTINEL [label](https://example.com/a_(b) "title")',
    "autolink": "LJ_SENTINEL <https://example.com?q=a&b=c>",
    "image": 'LJ_SENTINEL ![alt text](image.png "title")',
    "html": 'LJ_SENTINEL <span data-x="1">inline</span>',
    "escapes": r"LJ_SENTINEL \*literal\* \[brackets\] \\ slash",
    "entities": "LJ_SENTINEL &copy; &#x1F642; &MadeUpEntity;",
    "breaks": "LJ_SENTINEL first  \nsecond\\\nthird",
    "math": r"LJ_SENTINEL $x^2 + y^2$ and \(z^2\)",
    "nul": "LJ_SENTINEL before\x00after",
    "bidi": "LJ_SENTINEL abc \u202etext\u202c end",
    "long": f"LJ_SENTINEL {'word ' * 400}",
}

MARKDOWN_CASES = [
    pytest.param(block, inline, id=f"{block_name}-{inline_name}")
    for (block_name, block), (inline_name, inline) in product(
        MARKDOWN_BLOCK_CASES.items(), MARKDOWN_INLINE_CASES.items()
    )
]


def _walk_sections(section: SectionNode):
    yield section
    for child in section.children:
        yield from _walk_sections(child)


def _walk_blocks(blocks: list[DocumentBlock] | tuple[DocumentBlock, ...]):
    for block in blocks:
        yield block
        yield from _walk_blocks(block.children)


def _assert_tree_invariants(tree: DocTree) -> None:
    assert tree.root.level == 0
    assert tree.root.title == tree.title

    for section in _walk_sections(tree.root):
        for index, child in enumerate(section.children):
            assert child.index == index
            assert child.level > section.level
            assert child.path == (*section.path, (child.level, child.title))
        for block in _walk_blocks(section.blocks):
            assert block.text
            assert (block.start_line is None) == (block.end_line is None)
            if block.start_line is not None and block.end_line is not None:
                assert 1 <= block.start_line <= block.end_line


@pytest.mark.parametrize(("block_template", "inline"), MARKDOWN_CASES)
def test_markdown_parser_handles_combinatorial_syntax_corpus(
    block_template: str,
    inline: str,
) -> None:
    source = block_template.format(inline=inline)

    tree = MarkdownParser().parse(source, document_title="generated.md")

    _assert_tree_invariants(tree)
    rendered = "\n".join(
        [section.title for section in _walk_sections(tree.root)]
        + [
            block.text
            for section in _walk_sections(tree.root)
            for block in _walk_blocks(section.blocks)
        ]
    )
    assert "LJ_SENTINEL" in rendered


def test_markdown_parser_handles_deep_mixed_containers() -> None:
    source = "\n".join(
        [f"{'  ' * depth}- level-{depth}" for depth in range(24)]
        + ["", "> " * 20 + "LJ_SENTINEL"]
    )

    tree = MarkdownParser().parse(source, document_title="deep.md")

    _assert_tree_invariants(tree)
    assert "LJ_SENTINEL" in tree.root.blocks[-1].text


def _docx_bytes(build) -> bytes:
    from docx import Document

    document = Document()
    build(document)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _rewrite_docx(payload: bytes, transform) -> bytes:
    output = BytesIO()
    with (
        ZipFile(BytesIO(payload)) as source,
        ZipFile(output, "w", compression=ZIP_DEFLATED) as target,
    ):
        for item in source.infolist():
            target.writestr(item.filename, transform(item.filename, source.read(item)))
    return output.getvalue()


def _tree_text(tree: DocTree) -> str:
    return "\n".join(
        [section.title for section in _walk_sections(tree.root)]
        + [
            block.text
            for section in _walk_sections(tree.root)
            for block in _walk_blocks(section.blocks)
        ]
    )


def test_docx_parser_groups_consecutive_lists_and_preserves_body_order() -> None:
    def build(document) -> None:
        document.add_heading("Section", level=1)
        document.add_paragraph("bullet one", style="List Bullet")
        document.add_paragraph("bullet two", style="List Bullet")
        document.add_paragraph("between")
        document.add_paragraph("number one", style="List Number")
        document.add_paragraph("number two", style="List Number")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "left"
        table.cell(0, 1).text = "right"

    tree = DocxParser().parse(_docx_bytes(build), document_title="lists.docx")
    section = tree.root.children[0]

    assert [block.kind for block in section.blocks] == [
        "list",
        "paragraph",
        "list",
        "table",
    ]
    assert [len(section.blocks[index].children) for index in (0, 2)] == [2, 2]
    assert section.blocks[0].start_line is None
    assert [location.element_id for location in section.blocks[0].source_locations] == [
        "/w:document/w:body/w:p[2]",
        "/w:document/w:body/w:p[3]",
    ]
    assert [location.element_id for location in section.blocks[2].source_locations] == [
        "/w:document/w:body/w:p[5]",
        "/w:document/w:body/w:p[6]",
    ]
    assert section.blocks[0].attrs["ordered"] is False
    assert section.blocks[2].attrs["ordered"] is True
    _assert_tree_invariants(tree)


def test_docx_parser_detects_numbering_on_normal_paragraphs() -> None:
    from copy import deepcopy

    def build(document) -> None:
        for text, style_name in (
            ("direct bullet", "List Bullet"),
            ("direct number", "List Number"),
        ):
            paragraph = document.add_paragraph(text)
            style_num_pr = document.styles[style_name].element.pPr.numPr
            paragraph._p.get_or_add_pPr().append(deepcopy(style_num_pr))

    tree = DocxParser().parse(_docx_bytes(build), document_title="numbering.docx")

    assert [block.kind for block in tree.root.blocks] == ["list", "list"]
    assert [block.attrs["ordered"] for block in tree.root.blocks] == [False, True]


def test_docx_parser_does_not_infer_semantics_from_style_names_or_fonts() -> None:
    from docx.enum.style import WD_STYLE_TYPE

    def build(document) -> None:
        for style_name in ("Heading 2.1", "标题 1", "3.1.1", "List Bullet Custom"):
            document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            document.add_paragraph(f"text using {style_name}", style=style_name)
        quote = document.add_paragraph("style name alone is not a quote", style="Quote")
        quote.runs[0].font.name = "Consolas"

    tree = DocxParser().parse(_docx_bytes(build), document_title="styles.docx")

    assert tree.root.children == []
    assert [block.kind for block in tree.root.blocks] == ["paragraph"] * 5
    assert all(
        inline.kind == "text" for block in tree.root.blocks for inline in block.inlines
    )


def test_docx_parser_uses_explicit_outline_level_with_arbitrary_style_name() -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def build(document) -> None:
        style = document.styles.add_style(
            "Arbitrary semantic style", WD_STYLE_TYPE.PARAGRAPH
        )
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "2")
        style.element.get_or_add_pPr().append(outline)
        document.add_paragraph("Explicit level three", style=style)

    tree = DocxParser().parse(_docx_bytes(build), document_title="outline.docx")

    assert [(section.level, section.title) for section in tree.root.children] == [
        (3, "Explicit level three")
    ]


def test_docx_parser_keeps_distinct_numbering_instances_separate() -> None:
    def build(document) -> None:
        document.add_paragraph("first identity", style="List Bullet")
        document.add_paragraph("second identity", style="List Bullet 2")

    tree = DocxParser().parse(_docx_bytes(build), document_title="lists.docx")

    assert [block.kind for block in tree.root.blocks] == ["list", "list"]
    assert tree.root.blocks[0].attrs["num_id"] != tree.root.blocks[1].attrs["num_id"]
    assert [block.attrs["num_format"] for block in tree.root.blocks] == [
        "bullet",
        "bullet",
    ]


def test_docx_parser_rejects_unresolved_explicit_numbering() -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def build(document) -> None:
        paragraph = document.add_paragraph("declared but unresolved numbering")
        num_pr = OxmlElement("w:numPr")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "999")
        num_pr.append(num_id)
        paragraph._p.get_or_add_pPr().append(num_pr)

    with pytest.raises(ValueError, match="numId 999 resolves to 0"):
        DocxParser().parse(_docx_bytes(build), document_title="numbering.docx")


def test_docx_parser_omits_empty_list_paragraphs() -> None:
    def build(document) -> None:
        document.add_paragraph("", style="List Bullet")
        document.add_paragraph("visible item", style="List Bullet")

    tree = DocxParser().parse(_docx_bytes(build), document_title="empty-list.docx")

    assert len(tree.root.blocks) == 1
    assert [child.text for child in tree.root.blocks[0].children] == ["visible item"]


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


ONE_PIXEL_PNG = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUB"
    "AScY42YAAAAASUVORK5CYII="
)


def test_docx_parser_preserves_hyperlinks_images_and_inline_formatting() -> None:
    def build(document) -> None:
        paragraph = document.add_paragraph()
        paragraph.add_run("before ").bold = True
        _add_hyperlink(paragraph, "example", "https://example.com/path")
        paragraph.add_run(" after ").italic = True
        paragraph.add_run().add_picture(BytesIO(ONE_PIXEL_PNG))

    tree = DocxParser().parse(_docx_bytes(build), document_title="media.docx")
    paragraph = tree.root.blocks[0]

    assert [inline.kind for inline in paragraph.inlines] == [
        "strong",
        "link",
        "emphasis",
        "image",
    ]
    assert "[example](https://example.com/path)" in paragraph.text
    assert "![Picture 1](word/media/image1.png)" in paragraph.text
    assert paragraph.inlines[1].attrs["destination"] == "https://example.com/path"
    assert paragraph.inlines[3].attrs["destination"] == "word/media/image1.png"


def test_docx_parser_escapes_markdown_table_delimiters_and_line_breaks() -> None:
    def build(document) -> None:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "key|name"
        table.cell(0, 1).text = "path\\name"
        table.cell(1, 0).text = "first\nsecond"
        table.cell(1, 1).text = "value"

    tree = DocxParser().parse(_docx_bytes(build), document_title="table.docx")
    table = tree.root.blocks[0]

    assert table.text == (
        "| key\\|name | path\\\\name |\n| --- | --- |\n| first<br>second | value |"
    )


def test_docx_parser_preserves_links_and_images_inside_table_cells() -> None:
    def build(document) -> None:
        table = document.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        _add_hyperlink(paragraph, "example", "https://example.com")
        paragraph.add_run().add_picture(BytesIO(ONE_PIXEL_PNG))

    tree = DocxParser().parse(_docx_bytes(build), document_title="table-media.docx")

    assert "[example](https://example.com)" in tree.root.blocks[0].text
    assert "![Picture 1](word/media/image1.png)" in tree.root.blocks[0].text


def test_docx_parser_reads_visible_content_inside_ooxml_wrappers() -> None:
    from docx.oxml import OxmlElement

    def wrap(paragraph, wrapper_name: str, content_name: str | None = None) -> None:
        body = paragraph._p.getparent()
        position = body.index(paragraph._p)
        wrapper = OxmlElement(wrapper_name)
        content = OxmlElement(content_name) if content_name is not None else wrapper
        content.append(paragraph._p)
        if content is not wrapper:
            wrapper.append(content)
        body.insert(position, wrapper)

    def build(document) -> None:
        wrapped = document.add_paragraph("inside content control")
        inserted = document.add_paragraph("inside tracked insertion")
        deleted = document.add_paragraph("deleted text")
        wrap(wrapped, "w:sdt", "w:sdtContent")
        wrap(inserted, "w:ins")
        wrap(deleted, "w:del")

    tree = DocxParser().parse(_docx_bytes(build), document_title="wrappers.docx")
    rendered = _tree_text(tree)

    assert "inside content control" in rendered
    assert "inside tracked insertion" in rendered
    assert "deleted text" not in rendered
    assert [block.source_locations[0].element_id for block in tree.root.blocks] == [
        "/w:document/w:body/w:sdt/w:sdtContent/w:p",
        "/w:document/w:body/w:ins/w:p",
    ]


def test_docx_parser_uses_only_explicit_alternate_content_fallback() -> None:
    from docx.oxml import parse_xml

    def build(document) -> None:
        choice_paragraph = document.add_paragraph("unsupported choice")
        fallback_paragraph = document.add_paragraph("visible fallback")
        body = document.element.body
        alternate = parse_xml(
            '<mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/'
            'markup-compatibility/2006"><mc:Choice Requires="unsupported"/>'
            "<mc:Fallback/></mc:AlternateContent>"
        )
        choice, fallback = alternate
        choice.append(choice_paragraph._p)
        fallback.append(fallback_paragraph._p)
        body.insert(0, alternate)

    tree = DocxParser().parse(_docx_bytes(build), document_title="alternate.docx")

    assert [block.text for block in tree.root.blocks] == ["visible fallback"]


def test_docx_parser_ignores_xml_comments_between_body_elements() -> None:
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    def build(document) -> None:
        first = document.add_paragraph("before comment")
        document.add_paragraph("after comment")
        container = parse_xml(
            f"<w:body {nsdecls('w')}><!--producer annotation--></w:body>"
        )
        first._p.addnext(container[0])

    tree = DocxParser().parse(_docx_bytes(build), document_title="comments.docx")

    assert [block.text for block in tree.root.blocks] == [
        "before comment",
        "after comment",
    ]


def test_docx_parser_normalizes_strict_ooxml_namespaces() -> None:
    strict_uri_literal = "http://purl.oclc.org/ooxml/wordprocessingml/main"
    transitional = {
        b"http://schemas.openxmlformats.org/officeDocument/2006/relationships": (
            b"http://purl.oclc.org/ooxml/officeDocument/relationships"
        ),
        b"http://schemas.openxmlformats.org/wordprocessingml/2006/main": (
            b"http://purl.oclc.org/ooxml/wordprocessingml/main"
        ),
    }

    def transform(name: str, payload: bytes) -> bytes:
        if not name.endswith((".xml", ".rels")):
            return payload
        for standard, strict in transitional.items():
            payload = payload.replace(standard, strict)
        return payload

    strict_payload = _rewrite_docx(
        _docx_bytes(
            lambda document: document.add_paragraph(
                f"strict content {strict_uri_literal}"
            )
        ),
        transform,
    )
    tree = DocxParser().parse(strict_payload, document_title="strict.docx")

    assert tree.root.blocks[0].text == f"strict content {strict_uri_literal}"
    assert any("Strict OOXML" in repair for repair in tree.metadata["docx_repairs"])


def test_docx_parser_repairs_missing_image_content_type() -> None:
    def build(document) -> None:
        document.add_paragraph().add_run().add_picture(BytesIO(ONE_PIXEL_PNG))

    def transform(name: str, payload: bytes) -> bytes:
        if name != "[Content_Types].xml":
            return payload
        return payload.replace(
            b'<Default Extension="png" ContentType="image/png"/>', b""
        )

    damaged = _rewrite_docx(_docx_bytes(build), transform)
    tree = DocxParser().parse(damaged, document_title="missing-content-type.docx")

    assert tree.root.blocks[0].inlines[0].kind == "image"
    assert any(
        "content types for png" in repair for repair in tree.metadata["docx_repairs"]
    )


def test_docx_parser_reads_wrapped_table_cells() -> None:
    from docx.oxml import OxmlElement

    def build(document) -> None:
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "wrapped cell"
        row = cell._tc.getparent()
        position = row.index(cell._tc)
        content_control = OxmlElement("w:sdt")
        content = OxmlElement("w:sdtContent")
        content.append(cell._tc)
        content_control.append(content)
        row.insert(position, content_control)

    tree = DocxParser().parse(_docx_bytes(build), document_title="wrapped-cell.docx")

    assert "wrapped cell" in tree.root.blocks[0].text


def test_docx_parser_recognizes_inline_and_display_omml() -> None:
    from docx.oxml import OxmlElement

    def math(parent, text: str, *, display: bool = False) -> None:
        container = OxmlElement("m:oMathPara" if display else "m:oMath")
        expression = OxmlElement("m:oMath") if display else container
        run = OxmlElement("m:r")
        value = OxmlElement("m:t")
        value.text = text
        run.append(value)
        expression.append(run)
        if display:
            container.append(expression)
        parent.append(container)

    def build(document) -> None:
        paragraph = document.add_paragraph("equation ")
        math(paragraph._p, "a+b=c")
        math(document.element.body, "x=1", display=True)

    tree = DocxParser().parse(_docx_bytes(build), document_title="math.docx")

    assert tree.root.blocks[0].inlines[-1].kind == "math_inline"
    assert tree.root.blocks[0].inlines[-1].attrs["syntax"] == "omml"
    assert tree.root.blocks[1].kind == "math_block"
    assert tree.root.blocks[1].attrs == {"literal": "x=1", "syntax": "omml"}


@pytest.mark.parametrize("seed", range(32))
def test_docx_parser_handles_deterministic_generated_corpus(seed: int) -> None:
    rng = random.Random(seed)
    expected: list[str] = []

    def build(document) -> None:
        document.core_properties.title = f"Generated {seed}"
        for index in range(18):
            sentinel = f"LJ_DOCX_{seed}_{index}"
            expected.append(sentinel)
            choice = rng.randrange(6)
            if choice == 0:
                document.add_heading(sentinel, level=rng.randrange(1, 5))
            elif choice == 1:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(sentinel).bold = bool(index % 2)
            elif choice == 2:
                document.add_paragraph(sentinel, style="List Number")
            elif choice == 3:
                document.add_paragraph(sentinel, style="Quote")
            elif choice == 4:
                table = document.add_table(rows=1, cols=2)
                table.cell(0, 0).text = sentinel
                table.cell(0, 1).text = "值|value"
            else:
                paragraph = document.add_paragraph()
                paragraph.add_run(f"中文 👩🏽‍💻 {sentinel} ").italic = True
                paragraph.add_run("tail").underline = True

    tree = DocxParser().parse(_docx_bytes(build))
    rendered = _tree_text(tree)

    _assert_tree_invariants(tree)
    assert tree.title == f"Generated {seed}"
    assert all(sentinel in rendered for sentinel in expected)


@pytest.mark.parametrize("payload", [b"", b"not a zip", b"PK\x03\x04truncated"])
def test_docx_parser_rejects_invalid_packages(payload: bytes) -> None:
    with pytest.raises(BadZipFile):
        DocxParser().parse(payload, document_title="invalid.docx")


@pytest.mark.parametrize(
    "part_name",
    [
        "../word/document.xml",
        "/word/document.xml",
        "%2e%2e/word/document.xml",
        r"word\document.xml",
    ],
)
def test_docx_parser_rejects_unsafe_package_part_names(part_name: str) -> None:
    payload = BytesIO()
    with ZipFile(payload, "w", compression=ZIP_DEFLATED) as package:
        package.writestr(part_name, b"content")

    with pytest.raises(ValueError, match="unsafe DOCX package part name"):
        DocxParser().parse(payload.getvalue(), document_title="unsafe.docx")


def test_docx_parser_rejects_duplicate_package_part_names() -> None:
    payload = BytesIO()
    with (
        pytest.warns(UserWarning, match="Duplicate name"),
        ZipFile(payload, "w", compression=ZIP_DEFLATED) as package,
    ):
        package.writestr("word/document.xml", b"first")
        package.writestr("word/document.xml", b"second")

    with pytest.raises(ValueError, match="duplicate part names"):
        DocxParser().parse(payload.getvalue(), document_title="duplicate.docx")


# ---------------------------------------------------------------------------
# HTML, records, spreadsheet, database, and source parsers


HTML_BLOCK_CASES = {
    "body-paragraph": "<html><body><p>{inline}</p></body></html>",
    "fragment-paragraph": "<p>{inline}</p>",
    "fragment-bare-text": "{inline}",
    "fragment-bare-inline": "before <strong>{inline}</strong> after",
    "heading": "<h2>Title</h2>\n<p>{inline}</p>",
    "deep-heading": "<h1>Parent</h1>\n<h4>Deep</h4>\n<p>{inline}</p>",
    "blockquote": "<blockquote>{inline}</blockquote>",
    "nested-blockquote": "<blockquote>outer<blockquote>{inline}</blockquote></blockquote>",
    "bullet-list": "<ul><li>{inline}</li><li>second item</li></ul>",
    "ordered-list": '<ol start="7"><li>{inline}</li><li>second item</li></ol>',
    "nested-list-in-item": "<ul><li>parent<ul><li>{inline}</li></ul></li></ul>",
    "bare-li": "<li>{inline}</li><li>second",
    "fence": "<pre><code>{inline}\nsecond line</code></pre>",
    "table": "<table><tr><th>key</th></tr><tr><td>{inline}</td></tr></table>",
    "definition-list": "<dl><dt>term</dt><dd>{inline}</dd></dl>",
    "div": '<div class="note">{inline}</div>',
    "script": "<p>visible</p>\n<script>var hidden = 1;</script>\n<p>{inline}</p>",
    "style": "<style>p {{ color: red; }}</style>\n<p>{inline}</p>",
    "entities": "<p>a &amp; b &lt; c &copy; {inline}</p>",
    "comment": "<!-- comment -->\n<p>{inline}</p>\n<!-- tail -->",
    "img-alt": '<p>lead <img src="x.png" alt="decorative"> {inline}</p>',
    "mixed-newlines": "<html>\r\n<body>\r\n<p>{inline}</p>\r\n</body>\r\n</html>",
}

HTML_INLINE_CASES = {
    "plain": "LJ_SENTINEL plain text",
    "unicode": "LJ_SENTINEL 中文 Ελληνικά العربية 👩🏽‍💻",
    "strong": "LJ_SENTINEL <strong>bold</strong> text",
    "emphasis": "LJ_SENTINEL <em>italic</em> text",
    "code": "LJ_SENTINEL <code>code span</code> text",
    "link": 'LJ_SENTINEL <a href="https://example.com/a?b=c">label</a>',
    "image": 'LJ_SENTINEL <img src="image.png" alt="alt text">',
    "breaks": "LJ_SENTINEL first<br>second",
    "nested": "LJ_SENTINEL <strong>bold <em>both</em></strong>",
    "nul": "LJ_SENTINEL before\x00after",
    "bidi": "LJ_SENTINEL abc \u202etext\u202c end",
    "long": f"LJ_SENTINEL {'word ' * 200}",
}

HTML_CASES = [
    pytest.param(block, inline, id=f"{block_name}-{inline_name}")
    for (block_name, block), (inline_name, inline) in product(
        HTML_BLOCK_CASES.items(), HTML_INLINE_CASES.items()
    )
]


@pytest.mark.parametrize(("block_template", "inline"), HTML_CASES)
def test_html_parser_handles_combinatorial_syntax_corpus(
    block_template: str,
    inline: str,
) -> None:
    source = block_template.format(inline=inline)

    tree = HTMLParser().parse(source, document_title="generated.html")

    _assert_tree_invariants(tree)
    rendered = _tree_text(tree)
    assert "LJ_SENTINEL" in rendered


@pytest.mark.parametrize(
    ("source", "expected_tokens"),
    [
        ("<h1>Hello<h2>World", ("Hello", "World")),
        ("<!doctype html><h3><li>abc</h2>foo", ("abc", "foo")),
        ("<!doctype html><h1><div><h3><span></h1>foo", ("foo",)),
        (
            "<li>hello<li>world<ul>how<li>item</ul>you",
            ("hello", "world", "how", "item", "you"),
        ),
        ("<p>before</p>\n<table><tr><td>cell</td></tr>", ("before", "cell")),
        (
            "<ul><li>one<ul><li>nested</li></ul></li><li>two</li></ul>",
            ("one", "nested", "two"),
        ),
        ("<ul><li>one</li><ul><li>nested</li></ul></ul>", ("one", "nested")),
        ("<dl><dt>alpha</dt><dt>beta</dt></dl>", ("alpha", "beta")),
        ("text<div>more</div>tail", ("text", "more", "tail")),
        ("<p>unclosed paragraph at eof", ("unclosed", "eof")),
        ("<blockquote>unclosed quote", ("unclosed", "quote")),
        ("<ol><li>unclosed item", ("unclosed", "item")),
        ("<td>stray cell</td>", ("stray", "cell")),
        ("</p></li></ul></table>stray end tags", ("stray", "tags")),
        ("<h2>heading with <svg><circle/></svg> inline</h2>", ("heading", "inline")),
        (
            "<h2>heading with <button>btn</button> inline</h2>",
            ("heading", "btn", "inline"),
        ),
        ("<table><td>no tr</td><tr><td>row</td></tr></table>", ("row",)),
        ("<div><div><div><p>deep nesting</p></div></div></div>", ("deep", "nesting")),
    ],
    ids=lambda value: (
        value[:40].replace("/", "-").replace(" ", "_")[:40]
        if isinstance(value, str)
        else "tokens"
    ),
)
def test_html_parser_preserves_text_in_adversarial_fragments(
    source: str, expected_tokens: tuple[str, ...]
) -> None:
    tree = HTMLParser().parse(source, document_title="adversarial.html")

    _assert_tree_invariants(tree)
    extracted = _tree_text(tree)
    for token in expected_tokens:
        assert token in extracted, f"lost {token!r} in {source!r}"


def test_html_parser_applies_implied_end_tags_between_headings() -> None:
    tree = HTMLParser().parse("<h1>Hello<h2>World", document_title="implies.html")

    assert [(section.level, section.title) for section in tree.root.children] == [
        (1, "Hello")
    ]
    assert [(child.level, child.title) for child in tree.root.children[0].children] == [
        (2, "World")
    ]


def test_html_parser_renders_bare_list_items_as_implicit_list() -> None:
    tree = HTMLParser().parse(
        "<li>hello<li>world<ul>how<li>do</ul>you", document_title="bare-li.html"
    )

    assert [block.text for block in tree.root.blocks] == [
        "- hello\n- world",
        "how",
        "- do",
        "you",
    ]


def test_html_parser_flushes_unclosed_table_and_lists_at_eof() -> None:
    tree = HTMLParser().parse(
        "<p>before</p>\n<ul><li>item\n<table><tr><td>cell",
        document_title="unclosed.html",
    )

    assert [block.kind for block in tree.root.blocks] == [
        "paragraph",
        "list",
        "html_table",
    ]
    assert "cell" in tree.root.blocks[2].text
    _assert_tree_invariants(tree)


def test_html_parser_keeps_nested_list_text_inside_parent_item() -> None:
    tree = HTMLParser().parse(
        "<ul><li>one<ul><li>nested</li></ul></li><li>two</li></ul>",
        document_title="nested.html",
    )

    assert len(tree.root.blocks) == 1
    assert [child.text for child in tree.root.blocks[0].children] == [
        "one - nested",
        "two",
    ]


def test_html_parser_separates_block_level_containers() -> None:
    tree = HTMLParser().parse(
        "<dl><dt>alpha</dt><dt>beta</dt></dl>", document_title="dl.html"
    )

    assert "alpha beta" in _tree_text(tree).replace("\n", " ")


@pytest.mark.parametrize("seed", range(16))
def test_html_parser_handles_deterministic_generated_corpus(seed: int) -> None:
    rng = random.Random(seed)
    oracle = _RandomHtmlOracle(rng)
    source = oracle.build()

    tree = HTMLParser().parse(source, document_title=f"generated-{seed}.html")

    _assert_tree_invariants(tree)
    rendered = _tree_text(tree)
    assert all(sentinel in rendered for sentinel in oracle.sentinels)


class _RandomHtmlOracle:
    """Minimal independent HTML generator used by the seeded corpus test."""

    def __init__(self, rng: random.Random) -> None:
        self._rng = rng
        self._counter = 0
        self.sentinels: list[str] = []
        self._parts: list[str] = []

    def _sentinel(self) -> str:
        self._counter += 1
        sentinel = f"LJ_H_{self._counter:03d}"
        self.sentinels.append(sentinel)
        return sentinel

    def build(self) -> str:
        level = 1
        for _ in range(self._rng.randint(6, 18)):
            choice = self._rng.random()
            sentinel = self._sentinel()
            if choice < 0.2:
                level = min(6, max(1, level + self._rng.choice([-1, 1, 2])))
                self._parts.append(f"<h{level}>{sentinel} heading</h{level}>")
            elif choice < 0.5:
                self._parts.append(f"<p>{sentinel} paragraph text</p>")
            elif choice < 0.65:
                tag = "ol" if self._rng.random() < 0.5 else "ul"
                self._parts.append(
                    f"<{tag}><li>{sentinel} item one</li><li>item two</li></{tag}>"
                )
            elif choice < 0.8:
                self._parts.append(f"<table><tr><td>{sentinel} cell</td></tr></table>")
            else:
                self._parts.append(f"<blockquote>{sentinel} quote</blockquote>")
            if self._rng.random() < 0.2:
                self._parts.append("<!-- generated comment -->")
        return "\n".join(self._parts) + "\n"


@pytest.mark.parametrize("seed", range(8))
def test_records_parsers_handle_deterministic_generated_corpus(seed: int) -> None:
    from benchmarks.random_corpus import FORMAT_SPECS

    for name in ("csv", "tsv", "jsonl", "json", "yaml", "toml", "xml", "text", "log"):
        rng = random.Random(f"robustness:{seed}:{name}".encode())
        document = FORMAT_SPECS[name].generate(rng)
        if document.allowed_error_types:
            continue
        tree = (
            FORMAT_SPECS[name]
            .make_parser()
            .parse(document.source, document_title=document.document_id)
        )

        _assert_tree_invariants(tree)
        if document.required_elements and document.min_token_recall:
            assert _tree_text(tree)  # reference oracle is asserted in benchmarks


def test_delimited_text_parser_preserves_embedded_newlines_in_quoted_fields() -> None:
    tree = DelimitedTextParser().parse(
        'name,notes\nAda,"line one\nline two"\nGrace,ok', document_title="notes.csv"
    )

    assert [block.text for block in tree.root.blocks] == [
        "name: Ada\nnotes: line one\nline two",
        "name: Grace\nnotes: ok",
    ]


def test_delimited_text_parser_preserves_rfc4180_field_shapes() -> None:
    tree = DelimitedTextParser().parse(
        'a,b\n"x, y","z ""q"" w"\n,empty-first', document_title="rfc.csv"
    )

    assert [block.text for block in tree.root.blocks] == [
        'a: x, y\nb: z "q" w',
        "a: \nb: empty-first",
    ]


@pytest.mark.parametrize(
    ("parser_factory", "payload"),
    [
        (lambda: JSONParser(), "{not json"),
        (lambda: JSONParser(), ""),
        (lambda: JSONParser(), '["unterminated'),
        (lambda: YAMLParser(), "key: [unclosed"),
        (lambda: YAMLParser(), "\t- bad: indent"),
        (lambda: TOMLParser(), "key = "),
        (lambda: TOMLParser(), "[unclosed table"),
        (lambda: XMLParser(), "<root><unclosed></root>"),
        (lambda: XMLParser(), ""),
        (lambda: XMLParser(), "not xml at all"),
        (lambda: JSONLinesParser(), '{"valid": 1}\n{broken'),
        (lambda: JSONLinesParser(), "[1, 2,\n"),
        (lambda: NotebookParser(), "{not json"),
        (lambda: NotebookParser(), '{"cells": "not-a-list"}'),
    ],
    ids=lambda value: (
        value
        if isinstance(value, str)
        else value.__self__.__class__.__name__
        if hasattr(value, "__self__")
        else "case"
    ),
)
def test_records_parsers_reject_invalid_input_with_value_error(
    parser_factory, payload: str
) -> None:
    with pytest.raises(ValueError):
        parser_factory().parse(payload, document_title="invalid")


def test_xml_parser_retains_mixed_content_text_segments() -> None:
    tree = XMLParser().parse(
        "<root>lead in <b>bold</b> tail text<c>leaf</c></root>",
        document_title="mixed.xml",
    )

    assert [block.text for block in tree.root.blocks] == [
        "/root[1]/text()[1]: lead in",
        "/root[1]/b[1]: bold",
        "/root[1]/text()[2]: tail text",
        "/root[1]/c[1]: leaf",
    ]
    assert tree.root.blocks[0].attrs["text_segment"] is True


def test_xml_parser_counts_duplicate_sibling_elements() -> None:
    tree = XMLParser().parse("<r><a>1</a><a>2</a></r>", document_title="dup.xml")

    assert [block.text for block in tree.root.blocks] == [
        "/r[1]/a[1]: 1",
        "/r[1]/a[2]: 2",
    ]


def test_sql_parser_keeps_semicolons_inside_literals_and_comments() -> None:
    source = (
        "INSERT INTO t VALUES ('a;b');\n"
        "-- comment; with semicolon\n"
        "SELECT 'it''s';\n"
        "/* block; comment */ UPDATE t SET name = $$dollar; quoted$$;\n"
        'CREATE TABLE "weird;name" (id INTEGER)'
    )

    tree = SQLParser().parse(source, document_title="batch.sql")

    assert [block.text for block in tree.root.blocks] == [
        "INSERT INTO t VALUES ('a;b');",
        "-- comment; with semicolon\nSELECT 'it''s';",
        "/* block; comment */ UPDATE t SET name = $$dollar; quoted$$;",
        'CREATE TABLE "weird;name" (id INTEGER);',
    ]
    assert [
        (block.source_locations[0].line_start, block.source_locations[0].line_end)
        for block in tree.root.blocks
    ] == [(1, 1), (2, 3), (4, 4), (5, 5)]


def test_sql_parser_handles_empty_and_comment_only_input() -> None:
    tree = SQLParser().parse("-- only a comment\n", document_title="empty.sql")

    assert tree.root.blocks == []
    _assert_tree_invariants(SQLParser().parse("", document_title="empty.sql"))


@pytest.mark.parametrize("seed", range(4))
def test_tabular_parsers_handle_generated_workbooks_and_databases(seed: int) -> None:
    from benchmarks.random_corpus import FORMAT_SPECS

    for name in ("xlsx", "sqlite"):
        rng = random.Random(f"robustness:{seed}:{name}".encode())
        document = FORMAT_SPECS[name].generate(rng)
        if document.allowed_error_types:
            continue
        tree = (
            FORMAT_SPECS[name]
            .make_parser()
            .parse(document.source, document_title=document.document_id)
        )

        _assert_tree_invariants(tree)
        assert tree.root.blocks


def test_xlsx_parser_rejects_non_zip_payload() -> None:
    with pytest.raises(BadZipFile):
        XlsxParser().parse(b"not a zip archive", document_title="bad.xlsx")


def test_sqlite_parser_rejects_non_database_payload() -> None:
    import sqlite3

    with pytest.raises(sqlite3.DatabaseError):
        SQLiteParser().parse(b"definitely not sqlite", document_title="bad.db")


def test_xml_parser_rejects_dtd_and_entity_declarations() -> None:
    for payload in (
        '<!DOCTYPE root [<!ENTITY x "y">]><root>&x;</root>',
        "<root><!-- <!ENTITY x 'y'> --></root>",
    ):
        with pytest.raises(ValueError, match="DTD or entity"):
            XMLParser().parse(payload, document_title="hostile.xml")


def test_sqlite_parser_quotes_hostile_table_names() -> None:
    import sqlite3

    connection = sqlite3.connect(":memory:")
    connection.execute('CREATE TABLE "t"" --drop" (value TEXT)')
    connection.execute('INSERT INTO "t"" --drop" VALUES (\'kept\')')
    serialize = getattr(connection, "serialize")  # noqa: B009
    payload = serialize()
    connection.close()

    tree = SQLiteParser().parse(payload, document_title="hostile.db")

    assert [block.text for block in tree.root.blocks] == ["value: kept"]
    assert tree.root.blocks[0].attrs["table"] == 't" --drop'


@pytest.mark.parametrize(
    "language",
    ["python", "javascript", "typescript"],
)
def test_source_code_parser_symbol_fallback_matches_tree_sitter_contract(
    language: str, monkeypatch: pytest.MonkeyPatch
) -> None:
    from typing import cast

    from lumberjack.parser.code import parser as code_parser_module
    from lumberjack.parser.code.tree_sitter import CodeLanguage

    monkeypatch.setattr(
        code_parser_module, "extract_top_level_symbols", lambda *_args: None
    )
    if language == "python":
        source = "def one():\n    return 1\n\n\nclass Two:\n    pass\n"
        expected = 2
    elif language == "javascript":
        source = "function one() { return 1; }\n\nclass Two {}\n\nconst three = 3;\n"
        expected = 3
    else:
        source = (
            "function one(): number { return 1; }\n\n"
            "interface Shape { label: string }\n\n"
            "const three = 3;\n"
        )
        expected = 3
    tree = SourceCodeParser(language=cast(CodeLanguage, language)).parse(
        source, document_title="module"
    )

    assert len(tree.root.blocks) == expected
    assert all(block.kind == "record" for block in tree.root.blocks)
